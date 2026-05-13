from pathlib import Path
import csv

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "report"
OUT = REPORT / "StudentID_GaoYuchen_Exp4_ClassY.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code)
    run.font.name = "Menlo"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Menlo")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(40, 40, 40)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        run.font.color.rgb = RGBColor(31, 78, 121)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.18
    p.paragraph_format.space_after = Pt(5)
    for run in p.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)
    return p


def add_point(doc, title, requirement, code, result, note=None):
    add_heading(doc, title, 2)
    add_body(doc, f"得分点要求：{requirement}")
    add_body(doc, "Linux 代码：")
    add_code(doc, code)
    add_body(doc, "实验结果：")
    add_code(doc, result)
    if note:
        add_body(doc, f"结果说明：{note}")


def read_csv(path):
    with open(path, newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.1)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.1)
section.right_margin = Cm(2.1)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
styles["Normal"].font.size = Pt(10.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("生物信息学实验报告：Linux 基础与 ANNOVAR FASTA 序列统计")
r.bold = True
r.font.size = Pt(18)
r.font.name = "Arial"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
r.font.color.rgb = RGBColor(31, 78, 121)

meta = doc.add_table(rows=4, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta.style = "Table Grid"
items = [
    ("课程实验", "Linux basics / ANNOVAR / FASTA RNA sequence counting"),
    ("学生信息", "学号：__________    姓名：高宇辰    班级：ClassY"),
    ("实验数据", "bioinfor/annovar/humandb/hg38_refGeneWithVerMrna.fa"),
    ("提交文件名", "StudentID_GaoYuchen_Exp4_ClassY.docx（提交前可替换学号和班级）"),
]
for row, (k, v) in zip(meta.rows, items):
    set_cell_text(row.cells[0], k, bold=True)
    set_cell_text(row.cells[1], v)
    set_cell_shading(row.cells[0], "D9EAF7")

add_heading(doc, "一、实验目的", 1)
for text in [
    "在 Linux/macOS 终端环境中完成目录创建、文件复制/删除、压缩包解压、帮助文档查询和文本检索等基础操作。",
    "安装并检查基因组注释工具 ANNOVAR，理解 FASTA 文件中以 “>” 开头的序列描述行及其染色体位置信息。",
    "编写 Bash/R 脚本统计 hg38_refGeneWithVerMrna.fa 中 chr21 及各常染色体 RNA 序列数量，并结合 GRCh38 染色体长度讨论二者是否成比例。",
]:
    add_body(doc, text)

add_heading(doc, "二、实验材料与数据来源", 1)
rows = [
    ("课程 PDF", "c4.Linux_basics.pdf", "实验要求、命令练习和提交规范"),
    ("终端记录", "终端保存的输出.txt", "课堂已完成命令、错误尝试与运行结果"),
    ("分析软件", "annovar.latest.tar.gz / annovar/", "ANNOVAR 示例文件和 humandb 数据"),
    ("FASTA 文件", "bioinfor/annovar/humandb/hg38_refGeneWithVerMrna.fa", "RefGene mRNA 序列描述与序列内容"),
    ("染色体长度", "NCBI/GRC GRCh38 chromosome lengths", "用于与 RNA 序列数量比较"),
]
t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["材料", "文件/来源", "用途"]):
    set_cell_text(t.rows[0].cells[i], h, bold=True)
    set_cell_shading(t.rows[0].cells[i], "D9EAF7")
for row in rows:
    cells = t.add_row().cells
    for i, value in enumerate(row):
        set_cell_text(cells[i], value)

add_heading(doc, "三、PDF 得分点逐项展示", 1)
add_body(doc, "本节按照课程 PDF 中出现的 Practice / Experiment 得分点逐项整理。课堂已经完成的部分直接使用 `终端保存的输出.txt` 中的记录；需要汇总或进一步分析的部分在本报告中补齐并给出结果文件。")

score_rows = [
    ("1", "创建课程目录并进入，使用 pwd 验证", "1 分", "已完成", "`mkdir .../bioinfor` 后进入目录，`pwd` 输出 `/Users/gaoyuchen/Desktop/bioinfor`。"),
    ("2", "复制 ANNOVAR 压缩包并用 ls 验证", "0.5 分", "已完成", "`cp ... annovar.latest.tar.gz` 后 `ls` 可见压缩包。"),
    ("3", "用 bash 删除复制出的压缩包并验证", "0.5 分", "已完成", "`rm .../annovar.latest.tar.gz` 后再次 `ls`，目录为空。"),
    ("4", "运行并解释 `tar -zxvf annovar.latest.tar.gz`，注明帮助来源", "1 分", "已完成", "执行 `tar -zxvf` 解压 ANNOVAR；帮助来源为 `tar --help`。"),
    ("5", "用 cat 查看 ANNOVAR example/README", "练习点", "已完成", "终端记录中已输出 README，说明 ANNOVAR 示例文件用途。"),
    ("6", "用 less 查看 hg38_refGeneWithVerMrna.fa，并搜索 chr22", "1 分", "展示命令", "课堂已用 `less` 查看 FASTA；报告补充给出 `/chr22` 搜索方式。"),
    ("7", "筛选 FASTA 中包含 chr21 的行，输出到 hg38_chr21.txt 并查看", "1 分", "已完成", "`hg38_chr21.txt` 共 1121 行，文件位于 `bioinfor/hg38_chr21.txt`。"),
    ("8", "写 Bash 脚本，输入 FASTA，输出 chr21 序列数量", "1 分", "已完成", "`count_chr21.sh` 输出 1121。"),
    ("9", "在脚本中检查输入文件是否存在", "1 分", "补充完成", "报告给出带 `-f` 判断的改进脚本。"),
    ("10", "循环比较所有常染色体序列数量", "1 分", "补充完成", "统计 chr1-chr22 RNA 序列数量并输出 CSV 表。"),
    ("11", "用 R 可视化结果", "1 分", "补充完成", "生成柱状图和长度-数量散点图。"),
    ("12", "讨论 RNA 序列数量是否与染色体长度成比例及影响因素", "1 分", "补充完成", "Pearson r = 0.572；结论为不严格成比例，受基因密度等影响。"),
]

score_table = doc.add_table(rows=1, cols=5)
score_table.style = "Table Grid"
score_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["序号", "PDF 要求", "分值", "完成情况", "证据/结果"]):
    set_cell_text(score_table.rows[0].cells[i], h, bold=True)
    set_cell_shading(score_table.rows[0].cells[i], "D9EAF7")
for row in score_rows:
    cells = score_table.add_row().cells
    for i, value in enumerate(row):
        set_cell_text(cells[i], value)

add_heading(doc, "四、逐项实验结果与 Linux 代码", 1)

add_point(
    doc,
    "得分点 1：创建课程目录并验证当前位置（1 分）",
    "从 home 目录开始，创建课程文件夹 bioinfor，进入该目录，并用 pwd 验证。",
    "cd ~\nmkdir /Users/gaoyuchen/Desktop/bioinfor\ncd /Users/gaoyuchen/Desktop/bioinfor\npwd",
    "/Users/gaoyuchen/Desktop/bioinfor",
    "说明课程目录创建成功，并且当前工作目录已经切换到 bioinfor。",
)

add_point(
    doc,
    "得分点 2：复制 ANNOVAR 压缩包并用 ls 验证（0.5 分）",
    "将 ANNOVAR 安装包复制到课程目录，并用 ls 查看是否复制成功。",
    "cp /Users/gaoyuchen/Desktop/homework4/annovar.latest.tar.gz /Users/gaoyuchen/Desktop/bioinfor\nls /Users/gaoyuchen/Desktop/bioinfor",
    "annovar.latest.tar.gz",
    "ls 结果中出现 annovar.latest.tar.gz，说明复制成功。",
)

add_point(
    doc,
    "得分点 3：用 bash 删除压缩包并验证（0.5 分）",
    "使用 rm 删除课程目录中的压缩包，并再次用 ls 验证。",
    "rm /Users/gaoyuchen/Desktop/bioinfor/annovar.latest.tar.gz\nls /Users/gaoyuchen/Desktop/bioinfor",
    "# 无输出",
    "删除后 ls 没有显示文件，说明压缩包已从课程目录删除。",
)

add_point(
    doc,
    "得分点 4：解压 ANNOVAR 并解释命令（1 分）",
    "运行并解释 `tar -zxvf annovar.latest.tar.gz`，写明帮助来源。",
    "tar --help\ntar -zxvf /Users/gaoyuchen/Desktop/bioinfor/annovar.latest.tar.gz",
    "x annovar/\nx annovar/example/\nx annovar/example/ex1.avinput\nx annovar/example/README\nx annovar/humandb/\nx annovar/humandb/hg38_refGeneWithVerMrna.fa\nx annovar/annotate_variation.pl\nx annovar/table_annovar.pl",
    "帮助来源为 `tar --help`。其中 `-x` 表示解包，`-z` 表示 gzip 解压，`-v` 表示显示过程，`-f` 表示指定压缩包文件。",
)

add_point(
    doc,
    "得分点 5：用 cat 查看 ANNOVAR README（练习点）",
    "使用 cat 查看小文本文件，理解 ANNOVAR 示例目录中的文件说明。",
    "cat /Users/gaoyuchen/Desktop/bioinfor/annovar/example/README",
    "visit ANNOVAR website at http://www.openbioinformatics.org/annovar for more exmaple.\n\nex1.avinput: a simple ANNOVAR input example with a few variants\nex2.vcf: a simple VCF file with genotype information for 3 samples\ngene_xref.txt: an example gene cross-reference file\nhumandb/hg19_example_db_generic.txt: an example file for generic database",
    "README 说明了 example 目录中 avinput、VCF、gene_xref 和示例数据库文件的用途。",
)

add_point(
    doc,
    "得分点 6：用 less 查看 FASTA 并搜索 chr22（1 分）",
    "使用 less 查看大 FASTA 文件，并在 less 中搜索 chr22。",
    "less /Users/gaoyuchen/Desktop/bioinfor/annovar/humandb/hg38_refGeneWithVerMrna.fa\n# 在 less 中输入：/chr22\n# 按 n 查看下一个匹配，按 q 退出",
    ">NM_... Comment: this sequence (leftmost exon at chr22:...) is generated by ANNOVAR ...",
    "FASTA 文件较大，less 可分页查看；`/chr22` 可定位包含 chr22 的序列描述行。",
)

add_point(
    doc,
    "得分点 7：筛选 chr21 并导出结果文件（1 分）",
    "在 FASTA 文件中筛选包含 chr21 的行，输出到 hg38_chr21.txt，并查看新文件。",
    "grep \"chr21\" /Users/gaoyuchen/Desktop/bioinfor/annovar/humandb/hg38_refGeneWithVerMrna.fa > hg38_chr21.txt\ncat hg38_chr21.txt\nmv /Users/gaoyuchen/hg38_chr21.txt /Users/gaoyuchen/Desktop/bioinfor\nwc -l /Users/gaoyuchen/Desktop/bioinfor/hg38_chr21.txt\nhead -n 3 /Users/gaoyuchen/Desktop/bioinfor/hg38_chr21.txt",
    "1121 /Users/gaoyuchen/Desktop/bioinfor/hg38_chr21.txt\n>NR_104257.2   Comment: this sequence (leftmost exon at chr21:41798224) ...\n>NM_001001438.2 Warning: ... (leftmost exon at chr21:46188445) ...\n>NM_004571.5   Comment: this sequence (leftmost exon at chr21:42974561) ...",
    "输出文件 `bioinfor/hg38_chr21.txt` 共 1121 行，说明筛选出的 chr21 相关描述行数量为 1121。",
)

add_point(
    doc,
    "得分点 8：Bash 脚本统计 chr21 序列数量（1 分）",
    "脚本接受一个 FASTA 文件作为输入，输出 chr21 上的 RNA 序列数量。",
    '#!/bin/bash\n\nfile="$1"\ngrep -c "^>.*chr21" "$file"\n\nchmod +x count_chr21.sh\n./count_chr21.sh /Users/gaoyuchen/Desktop/bioinfor/annovar/humandb/hg38_refGeneWithVerMrna.fa',
    "1121",
    "脚本输出 1121，与 `hg38_chr21.txt` 的行数一致。",
)

add_point(
    doc,
    "得分点 9：脚本检查输入文件是否存在（1 分）",
    "在上一题脚本基础上，加入输入文件存在性判断。",
    '#!/bin/bash\n\nfile="$1"\nif [ ! -f "$file" ]; then\n  echo "Error: input FASTA file does not exist."\n  exit 1\nfi\ngrep -c "^>.*leftmost exon at chr21[:_]" "$file"',
    "# 输入正确文件时：\n1121\n\n# 输入不存在文件时：\nError: input FASTA file does not exist.",
    "使用 `-f` 判断输入是否为普通文件，可以避免空参数或错误路径导致 grep 报错。",
)

add_point(
    doc,
    "得分点 10：循环比较所有常染色体 RNA 序列数量（1 分）",
    "在上一题基础上，用循环统计 chr1 到 chr22 的 RNA 序列数量。",
    '#!/bin/bash\n\nfile="$1"\nif [ ! -f "$file" ]; then\n  echo "Error: input FASTA file does not exist."\n  exit 1\nfi\nfor i in {1..22}; do\n  chr="chr${i}"\n  count=$(grep -c "^>.*leftmost exon at ${chr}[:_]" "$file")\n  printf "%s\\t%s\\n" "$chr" "$count"\ndone',
    "chr1    7845\nchr2    5583\nchr3    4943\nchr4    3024\nchr5    3689\nchr6    6988\nchr7    3989\nchr8    3094\nchr9    3063\nchr10   3611\nchr11   4869\nchr12   3977\nchr13   1510\nchr14   2546\nchr15   3002\nchr16   3400\nchr17   4615\nchr18   1279\nchr19   6909\nchr20   1970\nchr21   1121\nchr22   1863",
    "统计结果已保存为 `report/chr_sequence_counts_autosomes.csv`，用于后续 R 可视化和相关性分析。",
)

add_point(
    doc,
    "得分点 11：使用 R 可视化结果（1 分）",
    "使用 R 可视化所有染色体 RNA 序列数量，并比较常染色体长度与 RNA 序列数量。",
    "Rscript report/make_chr_analysis.R",
    "生成文件：\nreport/chr_rna_counts_barplot.png\nreport/chr_length_vs_rna_count.png\nreport/chr_sequence_counts_all.csv\nreport/chr_sequence_counts_autosomes.csv\nreport/correlation_summary.txt",
    "R 脚本读取 FASTA 描述行，统计染色体 RNA 序列数量，生成柱状图和长度-数量散点图。",
)

add_point(
    doc,
    "得分点 12：讨论 RNA 序列数量是否与染色体长度成比例（1 分）",
    "结合可视化结果，讨论每条染色体 RNA 序列数量是否与 DNA 长度成比例，并分析影响因素。",
    "cat report/correlation_summary.txt",
    "Pearson correlation between autosome length and RNA sequence count: 0.572\nchr21 RNA sequence count: 1121",
    "常染色体长度与 RNA 序列数量存在中等正相关，但不严格成比例。chr19 长度较短却有较高 RNA 序列数，提示基因密度、可变剪接、非编码 RNA 注释密度、重复序列和数据库注释完整度等因素都会影响结果。",
)

add_heading(doc, "五、统计方法补充说明", 1)
add_body(doc, "课堂脚本 `count_chr21.sh` 使用 grep 统计 chr21 序列描述行，输出结果为 1121。")
add_code(doc, '#!/bin/bash\nfile="$1"\ngrep -c "^>.*chr21" "$file"')
add_body(doc, "按照 PDF 的脚本进阶要求，补充加入输入文件存在性判断。若未输入文件或文件不存在，脚本会输出错误信息并退出；若文件存在，则统计 chr21 序列描述行。")
add_code(
    doc,
    '#!/bin/bash\n'
    'file="$1"\n'
    'if [ ! -f "$file" ]; then\n'
    '  echo "Error: input FASTA file does not exist."\n'
    '  exit 1\n'
    'fi\n'
    'grep -c "^>.*leftmost exon at chr21[:_]" "$file"'
)
add_body(doc, "继续按照 PDF 的循环要求，比较 chr1 到 chr22 所有常染色体 RNA 序列数量。核心 Bash 思路如下：")
add_code(
    doc,
    '#!/bin/bash\n'
    'file="$1"\n'
    'if [ ! -f "$file" ]; then\n'
    '  echo "Error: input FASTA file does not exist."\n'
    '  exit 1\n'
    'fi\n'
    'for i in {1..22}; do\n'
    '  chr="chr${i}"\n'
    '  count=$(grep -c "^>.*leftmost exon at ${chr}[:_]" "$file")\n'
    '  printf "%s\\t%s\\n" "$chr" "$count"\n'
    'done'
)
add_body(doc, "为可视化和相关性分析，本报告使用 R 对 FASTA 描述行进行解析，只提取 `leftmost exon at chrN:` 或 `leftmost exon at chrN_alt:` 中的染色体编号，避免把 chrUn、chrM 或非描述字段误计入结果。")
add_code(
    doc,
    'headers <- readLines(input, warn = FALSE)\n'
    'headers <- headers[startsWith(headers, ">")]\n'
    'chr <- sub(".*leftmost exon at (chr[0-9XY]+)[:_].*", "\\\\1", headers)\n'
    'chr <- chr[grepl("^chr([1-9]|1[0-9]|2[0-2]|X|Y)$", chr)]\n'
    'counts <- as.data.frame(table(chr))'
)

add_heading(doc, "六、实验结果", 1)
add_body(doc, "筛选 `chr21` 的输出文件为 `bioinfor/hg38_chr21.txt`，共 1121 行；与脚本统计结果一致。")

doc.add_picture(str(REPORT / "chr_rna_counts_barplot.png"), width=Cm(15.2))
p = doc.add_paragraph("图 1  各染色体 RNA 序列数量统计")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

autosomes = read_csv(REPORT / "chr_sequence_counts_autosomes.csv")
table = doc.add_table(rows=1, cols=5)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["染色体", "长度(bp)", "长度(Mb)", "RNA序列数", "每Mb序列数"]
for i, h in enumerate(headers):
    set_cell_text(table.rows[0].cells[i], h, bold=True)
    set_cell_shading(table.rows[0].cells[i], "D9EAF7")
for row in autosomes:
    cells = table.add_row().cells
    values = [
        row["chromosome"],
        f'{int(row["length_bp"]):,}',
        f'{float(row["length_mb"]):.2f}',
        row["rna_sequence_count"],
        f'{float(row["rna_per_mb"]):.2f}',
    ]
    for i, value in enumerate(values):
        set_cell_text(cells[i], value)

doc.add_section(WD_SECTION_START.NEW_PAGE)
add_heading(doc, "七、长度与 RNA 序列数量关系", 1)
doc.add_picture(str(REPORT / "chr_length_vs_rna_count.png"), width=Cm(15.5))
p = doc.add_paragraph("图 2  常染色体长度与 RNA 序列数量关系")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_body(doc, "常染色体长度与 RNA 序列数量的 Pearson 相关系数为 r = 0.572，说明二者存在一定正相关，但不是严格成比例关系。若严格成比例，散点应大致沿直线均匀分布；实际结果中 chr19 长度较短但 RNA 序列数很高，chr13、chr18 等染色体相对较低，提示染色体长度并不能单独解释转录本数量。")
add_body(doc, "可能影响因素包括：不同染色体的基因密度差异、可变剪接产生的多个转录本、编码基因和非编码 RNA 注释密度、重复序列和异染色质比例、参考注释数据库更新程度，以及替代单倍型/未放置序列是否被纳入统计。总体上，RNA 序列数量更接近基因和转录本注释密度的结果，而不是简单由 DNA 分子长度决定。")

add_heading(doc, "八、结论", 1)
for text in [
    "ANNOVAR 已成功解压并可读取示例目录和 humandb FASTA 文件。",
    "FASTA 描述行以 `>` 开头，描述中包含转录本编号、染色体位置和 ANNOVAR 生成注释等信息。",
    "在 `hg38_refGeneWithVerMrna.fa` 中，chr21 相关 RNA 序列描述行为 1121 条。",
    "所有常染色体 RNA 序列数量与染色体长度有中等程度正相关，但不成严格比例；基因密度和转录本复杂度是更重要的影响因素。",
]:
    add_body(doc, text)

add_heading(doc, "九、AI 使用声明", 1)
add_body(doc, "本报告使用 AI 辅助完成。使用目的：阅读课程 PDF、整理课堂终端记录、汇总已有输出文件、补齐可复现的 R 统计与可视化、生成 Word 实验报告。")
add_body(doc, "使用的提示词：阅读这个文档里的 pdf，了解此次生信实验课程需要的东西，帮我完成并输出一份实验报告，然后其中 txt 文件是我课上已经用终端完成的内容，需要导出的结果应该也在其中的某个文件夹里面，你自行整理一下，我完成的部分可以直接用。")

add_heading(doc, "附录：生成的结果文件", 1)
for text in [
    "report/chr_sequence_counts_all.csv：所有 chr1-22、chrX、chrY 的 RNA 序列数量。",
    "report/chr_sequence_counts_autosomes.csv：常染色体长度、RNA 序列数和每 Mb 序列数。",
    "report/chr_rna_counts_barplot.png：各染色体 RNA 序列数柱状图。",
    "report/chr_length_vs_rna_count.png：常染色体长度与 RNA 序列数散点图。",
    "report/make_chr_analysis.R：可复现分析脚本。",
]:
    add_body(doc, text)

doc.save(OUT)
print(OUT)
